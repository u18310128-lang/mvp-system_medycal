from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"C:\Users\Administrador\Dev\tesis-sistema-citas\Borrador_metodologia_implementacion.docx"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1.18)
sec.right_margin = Inches(1.18)

styles = doc.styles
styles['Normal'].font.name = 'Arial'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
styles['Normal'].font.size = Pt(11)
styles['Normal'].paragraph_format.line_spacing = 1.5
styles['Normal'].paragraph_format.space_after = Pt(6)
for name, size in [('Heading 1', 14), ('Heading 2', 13), ('Heading 3', 12)]:
    st = styles[name]
    st.font.name = 'Arial'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    st.font.size = Pt(size)
    st.font.bold = True
    st.paragraph_format.space_before = Pt(12)
    st.paragraph_format.space_after = Pt(6)

def p(text='', style=None):
    para = doc.add_paragraph(style=style)
    para.add_run(text)
    return para

def bullet(text):
    para = doc.add_paragraph(style='List Bullet')
    para.add_run(text)
    return para

def table(rows, widths=None):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = 'Table Grid'
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.font.name = 'Arial'; run.font.size = Pt(9)
            if ridx == 0:
                for run in cells[i].paragraphs[0].runs:
                    run.bold = True
    return t

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('BORRADOR PARA EL CAPÍTULO III\nMETODOLOGÍA DE LA IMPLEMENTACIÓN DE LA SOLUCIÓN')
r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(14)
p('Texto preparado para adaptar a la tesis. Verificar fechas, responsables y evidencias antes de la entrega.', None)

p('3.3 Metodología de la implementación de la solución', 'Heading 1')
p('La implementación de la solución se organizó mediante una adaptación del marco de trabajo Scrum. Esta elección permitió desarrollar el sistema de forma progresiva, validar cada módulo y realizar ajustes conforme se conocían con mayor precisión las necesidades del consultorio. El desarrollo se orientó a construir un MVP compuesto por una PWA administrativa, un backend para la gestión de citas, una base de datos PostgreSQL, flujos de automatización en n8n y un agente conversacional conectado con WhatsApp.')

p('3.3.1 Selección de la metodología', 'Heading 2')
p('Para el desarrollo del MVP se seleccionó Scrum adaptado. La metodología resulta adecuada porque el sistema se compone de funcionalidades que pueden implementarse y validarse por separado, como la autenticación, la gestión de pacientes, la configuración de horarios, el registro de citas, los recordatorios y los indicadores administrativos.')
p('La aplicación de Scrum se ajustó al tamaño reducido del equipo del proyecto. En consecuencia, las funciones de Product Owner, Scrum Master y equipo de desarrollo fueron coordinadas por los responsables de la investigación. Cada iteración incluyó la planificación de actividades, el desarrollo de funcionalidades, las pruebas y la revisión de los resultados obtenidos.')
p('La metodología no se aplicó como un proceso aislado del trabajo de investigación. Los productos de cada iteración fueron contrastados con los requerimientos del consultorio, las reglas de negocio de la agenda y los indicadores definidos para la medición del ausentismo. De esta manera, el desarrollo técnico se mantuvo relacionado con los objetivos de la investigación.')

p('3.3.2 Desarrollo de la metodología', 'Heading 2')
p('El desarrollo se organizó en fases de inicio, planificación, implementación, revisión y cierre. Cada fase generó productos que sirvieron como entrada para la siguiente etapa. La siguiente descripción presenta las actividades realizadas y su relación con el MVP.')

p('3.3.2.1 Fase de inicio', 'Heading 3')
p('La fase de inicio permitió definir el problema, el propósito de la solución, los usuarios y los límites del proyecto. En esta etapa se estableció que el MVP sería una aplicación web progresiva (PWA), accesible desde computadoras, tablets y teléfonos móviles mediante un navegador, sin desarrollar aplicaciones nativas independientes para Android o iOS.')

p('Actividad 1: Elaboración de la visión del proyecto', 'Heading 3')
p('La visión se definió a partir de las dificultades observadas en el proceso manual de gestión de citas del Consultorio Perú Ruso. La confirmación, cancelación y reprogramación dependían del contacto realizado por el personal administrativo, lo que incrementaba el tiempo de trabajo y dificultaba el seguimiento de los pacientes.')
p('Como respuesta, se planteó desarrollar un sistema que centralice la agenda y automatice las comunicaciones relacionadas con las citas. El sistema permite administrar pacientes, médicos, horarios y citas desde una PWA. Además, utiliza n8n para coordinar los recordatorios y dispone de WhatsApp como canal conversacional principal, con correo electrónico como canal complementario cuando corresponda.')

p('Actividad 2: Elaboración del acta de constitución', 'Heading 3')
p('El acta de constitución estableció el objetivo, el alcance, los usuarios beneficiarios, las restricciones y los principales entregables. Se definieron tres roles operativos: recepción, médico y dirección. Recepción administra la agenda general; el médico consulta y actualiza únicamente la información relacionada con sus propias citas; y dirección supervisa las agendas, los usuarios, la auditoría y los indicadores.')
p('El alcance comprende el registro de pacientes, la configuración de horarios, la consulta de disponibilidad, el registro, confirmación, cancelación y reprogramación de citas, los recordatorios automatizados, el control de acceso por roles y la generación de indicadores. Quedan fuera la historia clínica electrónica, los diagnósticos, la teleconsulta, los pagos en línea, el envío de SMS y las llamadas automatizadas.')

p('Actividad 3: Determinación de los requerimientos del proyecto', 'Heading 3')
p('Los requerimientos se determinaron mediante la revisión del proceso actual y la identificación de las tareas realizadas por cada usuario. Se documentaron como historias de usuario y se priorizaron con la técnica MoSCoW. Las funciones clasificadas como Must correspondieron a las capacidades indispensables para operar el MVP; las funciones Should y Could se reservaron para mejoras posteriores.')
p('Entre los requerimientos principales se incluyeron el inicio de sesión, la gestión de pacientes, la configuración de horarios médicos, la prevención de citas superpuestas, la consulta de cupos, la reprogramación, la cancelación, el envío de recordatorios, la auditoría y la visualización de indicadores. También se estableció que el agente conversacional debe limitarse a consultar disponibilidad y gestionar operaciones de citas, sin brindar diagnósticos ni recomendaciones clínicas.')

p('Actividad 4: Recolección de información', 'Heading 3')
p('Se recopiló información sobre el proceso de registro de pacientes, la organización de los horarios, la duración de las atenciones, las excepciones de agenda, las confirmaciones y las reprogramaciones. También se identificaron los datos necesarios para registrar una cita y los eventos que deben conservarse en la bitácora de auditoría.')
p('La información obtenida permitió definir las entidades de la base de datos, las reglas de disponibilidad y los indicadores de evaluación. Asimismo, permitió establecer que los recordatorios no deben contener información clínica y que el contacto por WhatsApp o correo requiere un consentimiento previamente registrado.')

p('Actividad 5: Elaboración del diagrama del proceso', 'Heading 3')
p('Se elaboró el diagrama del proceso actual y del proceso propuesto. En el proceso actual, el personal registra las citas y realiza manualmente la confirmación con los pacientes. En el proceso propuesto, la cita se registra en el sistema, n8n programa los recordatorios y el paciente puede responder mediante WhatsApp.')
p('Para las solicitudes conversacionales, el flujo inicia en WhatsApp y continúa en n8n. Luego, el backend identifica al paciente, consulta las reglas y la disponibilidad en PostgreSQL, utiliza el modelo de lenguaje para interpretar la solicitud y devuelve una respuesta al canal de comunicación. La lógica de permisos y las reglas de las citas permanecen en el backend, mientras que n8n cumple la función de orquestar los mensajes y automatizaciones.')

p('3.3.2.2 Fase de planificación', 'Heading 3')
p('A partir de los requerimientos se elaboró el Product Backlog. Cada historia de usuario fue descrita con su objetivo, prioridad y criterio de aceptación. Posteriormente, las historias se agruparon en iteraciones de desarrollo considerando las dependencias técnicas y el valor que aportaban al funcionamiento del MVP.')
table([
    ['Sprint', 'Funcionalidades principales', 'Resultado esperado'],
    ['1', 'Autenticación, sesiones y roles', 'Acceso diferenciado para recepción, médico y dirección.'],
    ['2', 'Pacientes, médicos, horarios y cupos', 'Consulta de disponibilidad según horario real.'],
    ['3', 'Registro, cancelación y reprogramación', 'Gestión trazable del ciclo de la cita.'],
    ['4', 'n8n, WhatsApp y correo', 'Recordatorios y registro del estado de entrega.'],
    ['5', 'Agente conversacional', 'Consulta de cupos y operaciones de citas.'],
    ['6', 'Indicadores, auditoría, PWA y pruebas', 'MVP integrado y listo para validación.'],
])

p('3.3.2.3 Fase de implementación', 'Heading 3')
p('Durante la implementación se desarrollaron las funcionalidades priorizadas en el Product Backlog. El primer sprint se orientó a la autenticación y al control de permisos. En esta etapa se verificó que cada rol pudiera acceder únicamente a las operaciones autorizadas.')
p('En el segundo sprint se implementaron los módulos de pacientes, médicos, horarios y disponibilidad. La disponibilidad se calculó considerando los horarios configurados, las excepciones de agenda, las citas existentes y los cupos que ya habían sido ocupados.')
p('El tercer sprint incorporó el registro, la cancelación y la reprogramación de citas. Estas operaciones conservaron la trazabilidad de los cambios y validaron que una cita solo pudiera ser modificada por un usuario autorizado o por el paciente identificado en el canal conversacional.')
p('En el cuarto sprint se configuraron los flujos de recordatorios en n8n. El sistema programó los envíos en los hitos definidos para el proyecto y registró el resultado de cada intento. Cuando el envío por WhatsApp no pudo entregarse, se contempló el reintento y el uso del correo electrónico cuando existió una dirección válida y consentimiento registrado.')
p('El quinto sprint se destinó a la integración del agente conversacional. El agente se conectó con el backend mediante un endpoint protegido y utilizó herramientas para consultar disponibilidad, obtener las citas del paciente, registrar una nueva cita, confirmar, cancelar o reprogramar. Las decisiones sobre permisos y disponibilidad fueron aplicadas por el dominio y no por el modelo de lenguaje.')
p('En el sexto sprint se integraron los indicadores, la bitácora de auditoría y las características de la PWA. Finalmente, se verificó la instalación desde el navegador, la adaptación de la interfaz a distintos tamaños de pantalla y la disponibilidad de la aplicación cuando correspondió.')

p('3.3.2.4 Fase de revisión y validación', 'Heading 3')
p('Al finalizar cada sprint se revisaron las funcionalidades implementadas y se compararon con los criterios de aceptación definidos en el Product Backlog. Las pruebas incluyeron el inicio de sesión, la separación de permisos, la consulta de disponibilidad, la prevención de horarios superpuestos y las operaciones de registro, cancelación y reprogramación.')
p('El agente conversacional se validó mediante un banco de 31 casos. Se evaluó si identificaba correctamente la intención del paciente, si solicitaba los datos faltantes, si utilizaba la herramienta correspondiente y si respondía adecuadamente ante una solicitud fuera del alcance. También se registró la latencia de respuesta y el resultado de cada operación.')
p('En los flujos de n8n se verificó la activación de los recordatorios, la recepción de los eventos, el reintento ante fallos y el registro del estado del mensaje. Los errores detectados durante esta fase fueron documentados y corregidos antes de preparar el piloto.')

p('3.3.2.5 Fase de implementación y cierre', 'Heading 3')
p('Una vez completadas las iteraciones se realizó la integración de los módulos y se verificaron las conexiones con PostgreSQL, n8n, WhatsApp, el servicio de correo y el modelo de lenguaje. También se revisaron las variables de configuración, las credenciales y las reglas de acceso antes de utilizar el sistema en la etapa de evaluación.')
p('El cierre de la implementación comprendió la validación final del MVP, la preparación de los usuarios, la revisión de los registros de auditoría y la organización de los datos que serían utilizados en la comparación pretest y postest. De esta forma, la solución quedó preparada para medir los cambios en la gestión de citas y en el ausentismo de los pacientes.')

p('Nota de adaptación', 'Heading 2')
p('Antes de incorporar este borrador a la tesis, se deben reemplazar las expresiones que indiquen que una actividad fue realizada si todavía se encuentra pendiente. También se deben añadir las figuras, tablas, capturas y evidencias disponibles en el proyecto, especialmente el Product Backlog, el diagrama de proceso, los sprints, el flujo de n8n y las pruebas del agente.')

doc.save(OUT)
print(OUT)
